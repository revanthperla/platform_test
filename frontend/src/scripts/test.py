import os
import pandas as pd
from docx import Document

def extract_text_from_docx(docx_path, keywords):
    doc = Document(docx_path)
    extracted_data = {}

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for keyword in keywords:
            if keyword in text:
                before_keyword, after_keyword = text.split(keyword, 1)
                extracted_data[keyword] = after_keyword.strip()

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                for keyword in keywords:
                    if keyword in text:
                        before_keyword, after_keyword = text.split(keyword, 1)
                        extracted_data[keyword] = after_keyword.strip()

    return extracted_data

def process_forms(folder_path, output_file):
    keywords = ['Candidate name', 'Position applied for', 'Current location', 'Current employer', 'Total-experience', 'Total CTC', 'Expected CTC', 'Notice period (also if there is an option to buy-out)', 'Willingness to relocate with FAMILY – at the time of joining itself', 'Assessment Comments by HRINPUTS', 'Remarks']
    selected_keys = ['Candidate name', 'Position applied for', 'Current location', 'Current employer', 'Total-experience', 'Total CTC', 'Expected CTC', 'Notice period (also if there is an option to buy-out)', 'Willingness to relocate with FAMILY – at the time of joining itself', 'Assessment Comments by HRINPUTS', 'Remarks']

    all_extracted_data = []

    for file_name in os.listdir(folder_path):
        if file_name.endswith('.docx'):
            docx_path = os.path.join(folder_path, file_name)
            extracted_data = extract_text_from_docx(docx_path, keywords)
            all_extracted_data.append(extracted_data)

    # Combine all extracted data into a single DataFrame
    df = pd.DataFrame(all_extracted_data)

    # Transpose the DataFrame
    transposed_df = df.transpose()

    # Write DataFrame to Excel with formatting
    with pd.ExcelWriter(output_file, engine='xlsxwriter') as writer:
        transposed_df.to_excel(writer, index=False, header=False)

        # Get the xlsxwriter workbook and worksheet objects
        workbook = writer.book
        worksheet = writer.sheets['Sheet1']

        # Define a bold format
        bold_format = workbook.add_format({'bold': True})

        # Apply bold formatting to the keys (first row)
        for col_num, key in enumerate(selected_keys):
            worksheet.write(0, col_num, key, bold_format)


if __name__ == "__main__":
    folder_path = '/Users/apple/Documents'
    output_file = 'aggregated_data.xlsx'
    
    process_forms(folder_path, output_file)
